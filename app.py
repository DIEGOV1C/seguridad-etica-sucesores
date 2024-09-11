from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import os
import boto3
from datetime import datetime
from io import BytesIO
from docx.shared import Pt

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "https://acuerdo-sucesores.vercel.app"}})

# Configuración de S3 utilizando variables de entorno
S3_BUCKET = os.getenv("S3_BUCKET", "sucesores-data")
S3_KEY = os.getenv("S3_KEY")
S3_SECRET = os.getenv("S3_SECRET")
s3_client = boto3.client('s3', aws_access_key_id=S3_KEY, aws_secret_access_key=S3_SECRET)

def replace_text_in_table(table, replacements):
    """ Reemplaza texto en una tabla dentro del documento """
    for row in table.rows:
        for cell in row.cells:
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)

@app.route('/generate-agreement', methods=['POST'])
def generate_agreement():
    data = request.json

    # Descargar la plantilla desde S3
    try:
        s3_object = s3_client.get_object(Bucket=S3_BUCKET, Key='Acuerdo de seguridad.docx')
        doc_stream = BytesIO(s3_object['Body'].read())
        doc = Document(doc_stream)
    except Exception as e:
        return {"error": str(e)}, 500

    # Obtener la fecha actual
    current_date = datetime.now().strftime("%d/%m/%Y")

    # Reemplazar los marcadores de posición
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '<<COMPANY_NAME>>' in run.text:
                            run.text = run.text.replace('<<COMPANY_NAME>>', data['companyName'])
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                        if '<<REPRESENTATIVE_NAME>>' in run.text:
                            run.text = run.text.replace('<<REPRESENTATIVE_NAME>>', data['representativeName'])
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                        if '<<POSITION>>' in run.text:
                            run.text = run.text.replace('<<POSITION>>', data['position'])
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                        if '<<DATE>>' in run.text:
                            run.text = run.text.replace('<<DATE>>', current_date)
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)

    # Guardar el documento generado
    output_filename = f"Acuerdo_{data['companyName']}.docx"
    doc.save(output_filename)

    # Subir a S3
    try:
        s3_client.upload_file(output_filename, S3_BUCKET, output_filename)
    except Exception as e:
        return {"error": str(e)}, 500
    finally:
        # Eliminar el archivo local después de subir
        if os.path.exists(output_filename):
            os.remove(output_filename)

    return {"message": "El acuerdo ha sido almacenado correctamente."}, 200

@app.route('/health-check', methods=['GET'])
def health_check():
    return 'OK', 200


if __name__ == "__main__":
    # Configuración para entorno de producción
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)
